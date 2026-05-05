# -*- coding: utf-8 -*-
"""生成公文管理知识速记手册 Word 文档"""
import json, os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# === 数据 ===
data = {
    "title": "公文管理知识速记手册",
    "items": [
        # —— 公文行文方向 ——
        {
            "category": "公文行文方向",
            "content": "按照行文关系，公文可分为上行文、平行文、下行文三种。",
            "original": "按照行文关系可分为上行文件，平行文件。下行文件：",
            "corrections": ["文件→文", "添加顿号"]
        },
        {
            "category": "公文行文方向",
            "content": "上行文：指下级机关或业务部门向所属上级领导机关或业务主管部门的行文。如请示、报告。其行文方式有逐级行文、多级行文、越级行文三种。",
            "original": "上行文：指下级机关或业务部门向所属上级领导机关或业务主管部门的一种行文。如请示。报告。其行文方式有逐级行文，多级行文，越级行文三种。",
            "corrections": ["一种行文→的行文", "句号→顿号"]
        },
        {
            "category": "公文行文方向",
            "content": "请示、报告、议案三者统称为呈请性公文。",
            "original": "请示，报告，议案三者统称为呈请性公文",
            "corrections": ["逗号→顿号"]
        },
        {
            "category": "公文行文方向",
            "content": "下行文：指上级机关对所属下级机关的发文。如命令、指令、意见、决定、决议、布告、公告、通告、通知、通报、批复等。",
            "original": "下行文：指上级机关对所属下级机关的发文如命令。指令意见。决定决议，布告，公告通告通知通报，批复等",
            "corrections": ["发文如→发文。如", "句号→顿号（多处）"]
        },
        {
            "category": "公文行文方向",
            "content": "平行文：指同级机关，或者不相隶属的、没有领导与指导关系的机关、部门、单位之间的行文。主要是函，也包括一些通知、通报、纪要。",
            "original": "平行文：指同级机关，式者不相隶属的没有领导与指导关系的机关部门，单位之间的一种行文。主要是函。也包括一些通知，通报。纪要",
            "corrections": ["式者→或者", "句号→顿号（多处）"]
        },
        {
            "category": "公文行文方向",
            "content": "联系类公文是机关或单位之间相互商洽工作、询问和答复问题的公文，行文方向为平行文。主要的联系类公文有：函、纪要、通知、意见、通报。",
            "original": "联系类公文是机关或单位之间相互商治工作，询问和答复问题的公文，行文方向为平行文。主要的联系类公文有：幽，纪要通知，意见通报",
            "corrections": ["商治→商洽", "幽→函", "添加顿号"]
        },
        {
            "category": "公文行文方向",
            "content": "越级行文指下级机关在特殊情况下，越过直接上级向更高一级领导机关行文。以下几种特殊情况（一般不得越级行文）：\n（1）遇特殊重大紧急情况（如战争、天灾），如逐级上报会延误时机；\n（2）经多次请示直接上级，长期未得到解决的重大问题；\n（3）上级领导或领导机关交办，并指定越级上报的事项；\n（4）对直接上级机关或领导进行检举、控告；\n（5）直接上下级机关有争议而无法解决的重大问题；\n（6）询问、联系无需经过直接上级机关的一些工作问题等；\n（7）在市场经济中，为使文件精神尽快与群众见面，以使更好贯彻执行，采用电视、电话、广播、报刊等方式行文。",
            "original": "越级行文指下级机关在特殊情况下，越过直接上级向更高一级领导机关行文，下以几种特殊情况：(一般不得越级行文)...\n(1)遇特殊重大紧急，情况(战争。天定)如延级上报会延误时机\n(7)在市场经济中。为使文件精神尽快与群众见面，以使更好贯彻执行采用电视，电脑，电话。广播，报刊等方式行文",
            "corrections": ["下以→以下", "紧急，情况→紧急情况", "天定→天灾", "延级→逐级", "式→或（第4条）", "由→（删除多余字，第5条）", "电视，电脑→电视、电话（'电脑'疑似OCR误识别，公共行政语境通常为电视、电话、广播、报刊）"]
        },

        # —— 公文文种 ——
        {
            "category": "公文文种",
            "content": "通知：用于发布、传达要求下级机关执行和有关单位周知或执行的事项，批转、转发公文。",
            "original": "通知用于发布，传达要求下级机关执行和有关单位周知或执行的事项，批转，转发公文",
            "corrections": ["逗号→顿号"]
        },
        {
            "category": "公文文种",
            "content": "通告：适用于在一定范围内公布应当遵守或者周知的事项。",
            "original": "通告适用于在一定范围内公布应当遵守或国知的事项",
            "corrections": ["国知→周知"]
        },
        {
            "category": "公文文种",
            "content": "通报：适用于表彰先进、批评错误、传达重要精神和告知重要情况。",
            "original": "通报适用于表彰先进，批评错误。传达重要精神和告知重要情况",
            "corrections": ["逗号→顿号", "顿号→逗号（批评错误后）"]
        },
        {
            "category": "公文文种",
            "content": "决定：适用于对重要事项作出决策和部署、奖惩有关单位和人员、变更或者撤销下级机关不适当的决定事项。",
            "original": "决定适用于对重要事项作出决策和部署，奖惩有关单位和人员，变更或者撤销下级机关不适当的决定事项",
            "corrections": ["逗号→顿号（多处）"]
        },

        # —— 公文格式与效力 ——
        {
            "category": "公文格式与效力",
            "content": "公文被撤销，视为自始无效；公文被废止，视为自废止之日起失效。",
            "original": "公文被撤销视为自始无效，公文被废上，视身废止可起失\n效。",
            "corrections": ["废上→废止", "视身废止可起失→视为自废止之日起失", "合并为完整语句"]
        },
        {
            "category": "公文格式与效力",
            "content": "有特定发文机关标志的普发性公文和电报可以不加盖印章。有发文机关署名的，应当加盖发文机关印章。",
            "original": "有特定发文机关标志的普发性公文和电报可以不加盖印章，有\n发文机关署名的，应当加盖发文机关印章",
            "corrections": ["删除多余换行"]
        },
        {
            "category": "公文格式与效力",
            "content": "成文日期署会议通过或者发文机关负责人签发的日期。联合行文时，署最后签发机关负责人签发的日期。",
            "original": "成文日期署会议通过或者发文机关负责人签发的日期。联合行文时，署最后签发机关负责人签发的日期",
            "corrections": ["无错别字，仅删除标点异常"]
        },

        # —— 绩效评估 ——
        {
            "category": "政府绩效评估",
            "content": "在绩效评估体系中主要看绩效指标，而在政府部门的绩效指标中可以分四个方面：业绩、效率、效能、成本。",
            "original": "在绩效评估体系中主要看绩效指标而在政府部门的绩效指标中可以分四个方面：且绩。效率。效能。成本",
            "corrections": ["添加逗号", "且绩→业绩", "句号→顿号"]
        },
        {
            "category": "政府绩效评估",
            "content": "业绩指标通常包括：公共服务的数量和质量、公共管理目标的实现情况、政策制定水平与实施效果、公共管理的效益、公民对公共管理和公共服务的满意程度等。",
            "original": "此绩指标通常包括：公共服务的数量和质量。公共管理目标的实现情况，政策制定水平与实施效果，公共管理的效益，公民对公共管理和公共服务的满意程度等。",
            "corrections": ["此绩→业绩", "数量和质量。公共→数量和质量、公共"]
        },
        {
            "category": "政府绩效评估",
            "content": "效能指标包括行为的合理化水平、政府机关效能。",
            "original": "效能指标包括行为的合理化水平，政府机关效能",
            "corrections": ["逗号→顿号"]
        },
        {
            "category": "政府绩效评估",
            "content": "效率指标包括提供公共服务与产品的单位成本、服务与产品的数目、公共政策执行的开支、政府部门的办公物品损耗费用等。",
            "original": "效率指标包括提供公共服务与产品的单位成本，服务与产品的数目公共政策执行的开支。政府部门的办公物品损耗费用等。",
            "corrections": ["逗号→顿号", "数目公共→数目、公共", "开支。政府→开支、政府"]
        },
        {
            "category": "政府绩效评估",
            "content": "成本指标包括政府部门占用的人力、物力和财力，以及政府部门的支出。",
            "original": "成本指标包括政府部门占用的人力物力和财力。以及政府部门的支出",
            "corrections": ["人力物力和→人力、物力和", "。以及→、以及"]
        },

        # —— 沟通模式 ——
        {
            "category": "组织沟通模式",
            "content": "环式沟通特点：\n1. 封闭式控制结构，相当于链式两头相联结；\n2. 每个人同时可以与两侧沟通，地位平等；\n3. 组织的集中化程度和领导人预测程度较低，沟通速度较慢；\n4. 组织成员满意度高，适于创造高昂的士气。",
            "original": "环式沟通特点：1封闭式控制结构 相当子链连式两头相联结2每个人同时可以牙两侧沟地位平等3组织的集中化程度和领导人预测程度较低。打通速度较慢 4组织成员满意度高，适于创造高昂的士气",
            "corrections": ["1封闭→1.封闭", "相当子→相当于", "链连式→链式", "牙两侧沟→与两侧沟通", "3→3.", "。打通→，沟通", "4→4."]
        },

        # —— 公共服务 ——
        {
            "category": "公共服务",
            "content": "分众式公共服务：按不同标准区分群体，分众问需、分类施策，推行分众化特色服务，进一步满足多元服务需求，确保把民生落到实处。",
            "original": "分众式公共服务是按不同标准区分群体。分众问需。分类施策推行分众化特色服务进一步满足多元服务需求，确保把民生",
            "corrections": ["群体。分众→群体，分众", "问需。分类→问需、分类", "服务进一步→服务，进一步", "补充完整句尾（落到实处）"]
        },
    ]
}

# === 生成 Word ===
doc = Document()

# 设置默认字体
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 页边距
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# 标题
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run(data["title"])
run.bold = True
run.font.size = Pt(18)
run.font.name = '微软雅黑'
run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

doc.add_paragraph()  # 空行

# 按分类分组
from collections import OrderedDict
categories = OrderedDict()
for item in data["items"]:
    cat = item["category"]
    if cat not in categories:
        categories[cat] = []
    categories[cat].append(item)

# 颜色映射
cat_colors = {
    "公文行文方向": "1F4E79",
    "公文文种": "BF8F00",
    "公文格式与效力": "C00000",
    "政府绩效评估": "2E75B6",
    "组织沟通模式": "548235",
    "公共服务": "7030A0",
}

# 各分类标题颜色
header_colors = {
    "公文行文方向": RGBColor(0x1F, 0x4E, 0x79),
    "公文文种": RGBColor(0xBF, 0x8F, 0x00),
    "公文格式与效力": RGBColor(0xC0, 0x00, 0x00),
    "政府绩效评估": RGBColor(0x2E, 0x75, 0xB6),
    "组织沟通模式": RGBColor(0x54, 0x82, 0x35),
    "公共服务": RGBColor(0x70, 0x30, 0xA0),
}

idx = 1
for cat, items in categories.items():
    # 分类标题
    h = doc.add_paragraph()
    run = h.add_run(f"【{cat}】")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = header_colors.get(cat, RGBColor(0, 0, 0))
    run.font.name = '微软雅黑'
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for item in items:
        # 序号 + 内容
        lines = item["content"].split("\n")
        for i, line in enumerate(lines):
            p = doc.add_paragraph()
            if i == 0:
                prefix = f"{idx}. "
                idx += 1
            else:
                prefix = "    "
            run = p.add_run(prefix + line)
            run.font.size = Pt(11)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

        # 修改标注
        if item.get("corrections"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(4)
            corrs = "；".join(item["corrections"])
            run = p.add_run(f"  ✏️ {corrs}")
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
            run.font.name = '微软雅黑'
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# 统计摘要
doc.add_paragraph()
summary = doc.add_paragraph()
summary.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = summary.add_run("— 整理摘要 —")
run.bold = True
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

stats = [
    f"共整理 {len(data['items'])} 条知识点",
    f"分 {len(categories)} 个类别：{', '.join(categories.keys())}",
    "主要纠错类型：OCR错别字（幽→函、商治→商洽、且绩→业绩等）、标点符号误用（句号→顿号/逗号→顿号）、多余换行和乱码字",
]
for s in stats:
    p = doc.add_paragraph(s)
    for run in p.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

# 保存
output = os.path.join(os.path.dirname(os.path.abspath(__file__)), "公文管理知识速记手册.docx")
doc.save(output)
print(f"Done: {output}")
